# -*- coding: utf-8 -*-
"""Сборка прайса топ-20: заголовок, «Что это», «Что входит», таблица сложности (как в шаблоне).
   Запуск: python build_pricing_doc.py
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Dict, List, Sequence, Tuple

from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

SOURCE = r"C:\Users\vladi\Documents\ценообразование_по_услугам.docx"
# Новый простой файл (не перезаписывает старые варианты)
OUTPUT_DOCUMENTS = r"C:\Users\vladi\Documents\Прайс_услуг.docx"
OUTPUT_WORKSPACE = r"c:\Users\vladi\AgileBusinessVisitca\Прайс_услуг.docx"

TOP20_ORDERED: Tuple[int, ...] = (
    1,
    2,
    4,
    10,
    11,
    12,
    14,
    18,
    37,
    15,
    42,
    21,
    27,
    44,
    48,
    49,
    56,
    58,
    59,
    60,
)

MARKETING_SERVICE_NUMBERS = frozenset({58, 59, 60, 61, 65})

COLOR_MARKETING = RGBColor(0xB8, 0x45, 0x00)
COLOR_SECTION = RGBColor(0x1A, 0x4F, 0x8C)
COLOR_DEFAULT = RGBColor(0x00, 0x00, 0x00)


@dataclass(frozen=True)
class ServiceCopy:
    what: str
    includes: Tuple[str, ...]


# Тексты сути и состава — ориентир для клиента; цены ниже берутся только из исходного прайса
SERVICE_COPY: Dict[int, ServiceCopy] = {
    1: ServiceCopy(
        what="Независимая оценка того, насколько IT-инфраструктура, сервисы и процессы соответствуют задачам бизнеса, где есть риски и узкие места.",
        includes=(
            "сбор контекста: цели, текущий ландшафт систем, роли и ответственность",
            "проверка ключевых зон: серверы, сеть, резервирование, доступы, базовая кибергигиена",
            "анализ процессов эксплуатации и взаимодействия бизнеса с IT",
            "фиксация находок, приоритетов и рекомендаций в понятном отчёте",
        ),
    ),
    2: ServiceCopy(
        what="Документ и дорожная карта развития IT на горизонт 1–3 года: что строить, в каком порядке и с какой целью.",
        includes=(
            "выравнивание IT с бизнес-целями и ограничениями (бюджет, сроки, команда)",
            "целевая архитектура и принципы: что оставляем, что меняем, что внедряем",
            "этапы, приоритеты, зависимости между инициативами",
            "оценка рисков и метрик успеха (что считаем результатом)",
        ),
    ),
    4: ServiceCopy(
        what="Помогаем перевести процессы и сервисы в цифровой контур: от точечной автоматизации до программы изменений в компании.",
        includes=(
            "карта процессов «как сейчас» и гипотезы «как должно быть»",
            "подбор сценариев цифровизации: быстрые победы и стратегические шаги",
            "требования к системам, данным и интеграциям",
            "план внедрения, роли, обучение и сопровождение изменений",
        ),
    ),
    10: ServiceCopy(
        what="Одна посадочная страница под конкретное предложение: быстро объяснить ценность, снять возражения и привести к целевому действию.",
        includes=(
            "структура блоков и сценарий скролла (оффер, выгоды, соцдоказательства, FAQ, форма)",
            "адаптивная вёрстка и базовая скорость загрузки",
            "формы, кнопки, микроразметка по необходимости",
            "подключение аналитики и целей (по согласованию)",
        ),
    ),
    11: ServiceCopy(
        what="Сайт компании как цифровое представительство: услуги, кейсы, контакты, материалы для партнёров и клиентов.",
        includes=(
            "информационная архитектура и навигация",
            "дизайн и вёрстка основных типов страниц",
            "CMS для самостоятельного обновления контента (при выборе такого пути)",
            "формы обратной связи, базовая SEO-подготовка, интеграции по ТЗ",
        ),
    ),
    12: ServiceCopy(
        what="Онлайн-продажи: каталог, корзина, оформление заказа и связка с оплатой, доставкой и учётом на стороне бизнеса.",
        includes=(
            "каталог, карточка товара, фильтры и поиск (по объёму проекта)",
            "корзина, оформление, уведомления, личный кабинет при необходимости",
            "интеграции с оплатой и службами доставки (по согласованию)",
            "админка для контента и заказов, базовые отчёты",
        ),
    ),
    14: ServiceCopy(
        what="Веб-сервис с подпиской или ролями пользователей: ядро продукта в браузере, а не просто сайт-визитка.",
        includes=(
            "модель пользователей, роли и права доступа",
            "ключевые сценарии продукта (MVP или полный функционал)",
            "API и интеграции с внешними системами при необходимости",
            "базовый мониторинг, логирование, подготовка к нагрузке по ТЗ",
        ),
    ),
    18: ServiceCopy(
        what="Регулярное сопровождение живого сайта или сервиса: правки, обновления, мелкие задачи без отдельного проекта каждый раз.",
        includes=(
            "консультации и оценка задач в рамках пакета",
            "исправления, доработки интерфейса и логики по согласованному бэклогу",
            "обновления зависимостей/платформы в безопасных пределах",
            "реакция на инциденты согласно SLA (уровень зависит от договорённости)",
        ),
    ),
    37: ServiceCopy(
        what="Одно приложение сразу под iOS и Android на общей кодовой базе (например, Flutter/React Native): быстрее вывод на рынок при типовых сценариях.",
        includes=(
            "проектирование экранов и пользовательских потоков",
            "реализация клиента под обе платформы и публикация в сторах (по договорённости)",
            "работа с push, офлайн-режимом, гео — по ТЗ",
            "интеграция с backend/API и аналитикой",
        ),
    ),
    15: ServiceCopy(
        what="Веб-система для продаж, заказов, производства или учёта: единое место правды для отделов вместо таблиц и разрозненных сервисов.",
        includes=(
            "модель сущностей: клиенты, сделки, заказы, номенклатура и т.д.",
            "роли пользователей, воронки, статусы, базовая автоматизация",
            "отчёты и выгрузки, интеграции с сайтом, телефонией, складом — по ТЗ",
            "обучение ключевых пользователей и передача документации",
        ),
    ),
    42: ServiceCopy(
        what="Панели для руководителей и команд: наглядные показатели из ваших баз и сервисов без ручной «склейки» в Excel.",
        includes=(
            "согласование набора метрик и источников данных",
            "подготовка/очистка данных, модель для отчётов",
            "дашборды, фильтры, срезы, расписание обновления",
            "доступ по ролям и базовое сопровождение после запуска",
        ),
    ),
    21: ServiceCopy(
        what="Превращаем сырые данные в понятные выводы: качество данных, закономерности, основа для решений и следующих шагов (в т.ч. ML).",
        includes=(
            "инвентаризация источников и проверка полноты/качества",
            "очистка, нормализация, признаки для анализа",
            "разведочный анализ, гипотезы, визуализации",
            "итоговый отчёт с выводами и рекомендациями",
        ),
    ),
    27: ServiceCopy(
        what="Прикладной сервис с элементами ИИ: от прототипа до работающего API/интерфейса под задачу бизнеса.",
        includes=(
            "формализация задачи и критериев качества",
            "выбор подхода: готовые модели, дообучение, классические алгоритмы",
            "реализация backend/API и при необходимости UI",
            "базовый мониторинг, логирование запросов, план доработок",
        ),
    ),
    44: ServiceCopy(
        what="Связка вашего продукта или внутренней системы с внешним сервисом через API: обмен данными по расписанию или в реальном времени.",
        includes=(
            "разбор документации API и сценариев обмена",
            "реализация авторизации, запросов, обработки ответов и ошибок",
            "логирование, идемпотентность/повторы где нужно",
            "тесты на тестовом контуре, передача инструкции по эксплуатации",
        ),
    ),
    48: ServiceCopy(
        what="Имитация действий злоумышленника с разрешения заказчика: найти реальные векторы атаки до того, как ими воспользуются другие.",
        includes=(
            "согласование объёма (веб, API, сеть, социнженерия — по договорённости)",
            "сбор информации, сканирование, ручная проверка уязвимостей",
            "оформление находок: критичность, воспроизведение, рекомендации",
            "закрывающий отчёт для технических и управленческих ролей",
        ),
    ),
    49: ServiceCopy(
        what="Системная оценка зрелости ИБ: политики, доступы, конфигурации, процессы реагирования — не только «техника», но и организация.",
        includes=(
            "интервью и сбор артефактов (регламенты, схемы, списки систем)",
            "проверка ключевых контролей и типовых нарушений",
            "карта рисков и несоответствий нормам/лучшим практикам",
            "приоритизированный план работ с ориентирами по срокам",
        ),
    ),
    56: ServiceCopy(
        what="Корректный сбор данных о посещениях и действиях на сайте: чтобы цифры в отчётах совпадали с реальностью.",
        includes=(
            "настройка счётчиков и потоков данных",
            "цели, события, e-commerce (если применимо), фильтрация мусорного трафика",
            "базовые отчёты и проверка согласованности метрик",
            "рекомендации по доработке разметки и структуры данных",
        ),
    ),
    58: ServiceCopy(
        what="Оценка эффективности маркетинга: каналы, кампании, стоимость привлечения и вклад в результат.",
        includes=(
            "согласование модели атрибуции и набора KPI",
            "сбор данных из рекламных кабинетов и аналитики",
            "срезы по каналам/кампаниям/гео/продуктам",
            "выводы и приоритеты по оптимизации бюджета",
        ),
    ),
    59: ServiceCopy(
        what="Связываем рекламу, заявки, CRM и продажи: видим полный путь от клика до денег, а не обрывки в разных системах.",
        includes=(
            "проектирование сквозной схемы данных и идентификаторов",
            "настройка передачи событий и офлайн-конверсий",
            "контроль качества данных и расхождений между системами",
            "дашборды или регламентные отчёты для маркетинга и продаж",
        ),
    ),
    60: ServiceCopy(
        what="Рост доли посетителей, которые выполняют целевое действие: не «перекрасить кнопку», а системная работа с воронкой и гипотезами.",
        includes=(
            "аудит ключевых страниц и точек потерь",
            "формулировка гипотез, приоритетов и плана тестов",
            "подготовка макетов/правок и контроль внедрения",
            "измерение эффекта и накопление библиотеки улучшений",
        ),
    ),
}


@dataclass
class Level:
    complexity: str
    price: str
    description: str


@dataclass
class Service:
    num: int
    title: str
    direction: str
    levels: List[Level]


def parse_services(doc: Document) -> List[Service]:
    lines = [p.text.strip() for p in doc.paragraphs]
    services: List[Service] = []
    i = 0
    num_re = re.compile(r"^(\d+)\.\s+(.+)$")

    while i < len(lines):
        line = lines[i]
        m = num_re.match(line)
        if not m:
            i += 1
            continue
        num = int(m.group(1))
        title = m.group(2).strip()
        i += 1
        direction = ""
        if i < len(lines) and lines[i].startswith("Направление:"):
            direction = lines[i].replace("Направление:", "").strip()
            i += 1
        levels: List[Level] = []
        while i < len(lines):
            if num_re.match(lines[i]):
                break
            if lines[i] in ("Низкая сложность", "Средняя сложность", "Высокая сложность"):
                comp = lines[i]
                i += 1
                price = ""
                desc = ""
                if i < len(lines) and lines[i].startswith("Стоимость:"):
                    price = lines[i].replace("Стоимость:", "").strip()
                    i += 1
                if i < len(lines) and lines[i].startswith("Описание уровня:"):
                    desc = lines[i].replace("Описание уровня:", "").strip()
                    i += 1
                levels.append(Level(comp, price, desc))
            else:
                i += 1
        services.append(Service(num, title, direction, levels))
    services.sort(key=lambda s: s.num)
    return services


def _paragraph_list_bullet(doc: Document) -> Paragraph:
    try:
        return doc.add_paragraph(style="List Bullet")
    except KeyError:
        p = doc.add_paragraph()
        p.add_run("• ")
        return p


def _add_bullet_formatted(doc: Document, line: str) -> None:
    p = _paragraph_list_bullet(doc)
    p.paragraph_format.left_indent = Pt(18)
    if ": " in line:
        label, rest = line.split(": ", 1)
        lb = p.add_run(label + ": ")
        lb.bold = True
        lb.font.size = Pt(11)
        r = p.add_run(rest)
        r.font.size = Pt(11)
    else:
        r = p.add_run(line)
        r.font.size = Pt(11)


def _set_cell_run(cell, text: str, *, bold: bool = False, color: RGBColor | None = None, size_pt: float = 11) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _complexity_column_title(full: str) -> str:
    if full.endswith(" сложность"):
        return full[: -len(" сложность")]
    return full


def add_complexity_price_table(doc: Document, svc: Service, *, marketing: bool) -> None:
    """Таблица: строка заголовков сложности, строка описаний, строка стоимости."""
    levels = svc.levels
    if len(levels) != 3:
        raise ValueError(f"Услуга {svc.num}: ожидается 3 уровня сложности, получено {len(levels)}")

    accent = COLOR_MARKETING if marketing else COLOR_DEFAULT

    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(10)

    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        row.height = None
    col_widths = (Cm(3.2), Cm(4.7), Cm(4.7), Cm(4.7))
    for row in table.rows:
        for idx, w in enumerate(col_widths):
            row.cells[idx].width = w

    hdr = tuple(_complexity_column_title(lv.complexity) for lv in levels)
    _set_cell_run(table.rows[0].cells[0], "Сложность", bold=True, size_pt=11)
    for c in range(3):
        _set_cell_run(table.rows[0].cells[c + 1], hdr[c], bold=True, color=accent, size_pt=11)

    _set_cell_run(table.rows[1].cells[0], "Описание уровня", bold=True, size_pt=11)
    for c in range(3):
        _set_cell_run(table.rows[1].cells[c + 1], levels[c].description, bold=False, size_pt=11)

    _set_cell_run(table.rows[2].cells[0], "Стоимость", bold=True, size_pt=11)
    for c in range(3):
        _set_cell_run(table.rows[2].cells[c + 1], levels[c].price, bold=False, size_pt=11)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def write_service(doc: Document, svc: Service) -> None:
    marketing = svc.num in MARKETING_SERVICE_NUMBERS
    title_color = COLOR_MARKETING if marketing else COLOR_DEFAULT

    p0 = doc.add_paragraph()
    p0.paragraph_format.space_before = Pt(14)
    t = p0.add_run(f"{svc.num}. {svc.title}")
    t.bold = True
    t.font.size = Pt(14)
    t.font.color.rgb = title_color

    copy = SERVICE_COPY.get(svc.num)
    if copy:
        pw = doc.add_paragraph()
        w0 = pw.add_run("Что это: ")
        w0.bold = True
        w0.font.size = Pt(11)
        wr = pw.add_run(copy.what)
        wr.font.size = Pt(11)

        ph = doc.add_paragraph()
        h = ph.add_run("Что входит:")
        h.bold = True
        h.font.size = Pt(11)
        for line in copy.includes:
            _add_bullet_formatted(doc, line)
    else:
        pf = doc.add_paragraph()
        pf.add_run("Описание состава услуги уточняется по запросу.").italic = True

    add_complexity_price_table(doc, svc, marketing=marketing)


def add_section_heading(doc: Document, title: str) -> None:
    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(20)
    hp.paragraph_format.space_after = Pt(6)
    r = hp.add_run(title)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = COLOR_SECTION


def _collect_services(all_svcs: Dict[int, Service], nums: Sequence[int]) -> List[Service]:
    return [all_svcs[n] for n in nums if n in all_svcs]


def main() -> None:
    if len(TOP20_ORDERED) != 20:
        raise ValueError(f"Ожидается ровно 20 позиций, сейчас: {len(TOP20_ORDERED)}")
    if set(SERVICE_COPY.keys()) != set(TOP20_ORDERED):
        raise RuntimeError("SERVICE_COPY должен содержать ровно ключи TOP20")

    src = Document(SOURCE)
    all_svcs: Dict[int, Service] = {s.num: s for s in parse_services(src)}

    groups: List[Tuple[str, Tuple[int, ...]]] = [
        ("IT-стратегия и цифровая трансформация", (1, 2, 4)),
        ("Веб-разработка", (10, 11, 12, 14, 18)),
        ("Мобильные и корпоративные системы", (37, 15, 42)),
        ("Данные и AI", (21, 27)),
        ("Интеграции и безопасность", (44, 48, 49)),
        ("Аналитика и маркетинг", (56, 58, 59, 60)),
    ]

    flat = [n for _, nums in groups for n in nums]
    if sorted(flat) != sorted(TOP20_ORDERED):
        raise RuntimeError("Группы не совпадают с TOP20_ORDERED")

    out = Document()
    t = out.add_heading("Прайс услуг", level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = out.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = sub.add_run("20 позиций. Уровни и цены — из полного прайса. Маркетинг выделен цветом.")
    rr.italic = True
    rr.font.size = Pt(11)

    for section_title, nums in groups:
        add_section_heading(out, section_title)
        for n in nums:
            svc = all_svcs.get(n)
            if svc:
                write_service(out, svc)

    saved: List[str] = []
    for path in (OUTPUT_DOCUMENTS, OUTPUT_WORKSPACE):
        try:
            out.save(path)
            saved.append(path)
        except PermissionError:
            continue
    if not saved:
        raise PermissionError(
            "Не удалось сохранить (закройте Прайс_услуг.docx в Word и повторите запуск)."
        )
    for p in saved:
        print("Saved:", p)


if __name__ == "__main__":
    main()
